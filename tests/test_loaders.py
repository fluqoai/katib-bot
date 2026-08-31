"""Unit tests for rag/loaders.py — no live services needed."""

from __future__ import annotations

import zipfile
import shutil
from pathlib import Path

import pytest

from rag.loaders import SUPPORTED_EXTS, load_any, load_text


pytestmark = pytest.mark.unit


def test_supported_extensions():
    assert ".md" in SUPPORTED_EXTS
    assert ".docx" in SUPPORTED_EXTS
    assert ".pdf" in SUPPORTED_EXTS
    assert ".txt" in SUPPORTED_EXTS


def test_load_text(tmp_path: Path):
    f = tmp_path / "doc.md"
    f.write_text("# Hello\n\nمرحبا بالعالم", encoding="utf-8")
    assert load_text(f) == "# Hello\n\nمرحبا بالعالم"


def test_load_any_dispatches_text(tmp_path: Path):
    f = tmp_path / "doc.txt"
    f.write_text("hello", encoding="utf-8")
    assert load_any(f) == "hello"


def test_load_any_dispatches_markdown(tmp_path: Path):
    f = tmp_path / "doc.md"
    f.write_text("# Title", encoding="utf-8")
    assert load_any(f) == "# Title"


def test_load_any_raises_on_unsupported(tmp_path: Path):
    f = tmp_path / "image.png"
    f.write_bytes(b"\x89PNG")
    with pytest.raises(ValueError, match="Unsupported file extension"):
        load_any(f)


def test_load_docx_round_trip(tmp_path: Path):
    """Create a minimal valid .docx in memory and read it back."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()
    doc.add_heading("عنوان", level=1)
    doc.add_paragraph("هذا فقرة اختبار.")
    doc.add_paragraph("وفقرة ثانية فيها كلمة مهمة: سياسة الخصوصية.")
    p = tmp_path / "test.docx"
    doc.save(str(p))

    text = load_any(p)
    assert "عنوان" in text
    assert "فقرة اختبار" in text
    assert "سياسة الخصوصية" in text


def test_load_pdf_round_trip(tmp_path: Path):
    """Create a minimal valid .pdf and read it back.

    Note: PyMuPDF's default fonts (helv, tiro, etc.) are Latin-only.
    For Arabic PDF text extraction to work, the source PDF must embed
    a Unicode-aware font (e.g. Amiri, Noto Sans Arabic). The user's
    actual `لوائح_المركز_الوطني/*.pdf` files do, but a synthetic test
    PDF built with the default font won't round-trip Arabic cleanly —
    so we test the loader with Latin text here, and trust real-world
    Arabic PDFs (which we can't generate in a unit test) for Arabic.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        pytest.skip("PyMuPDF not installed")

    p = tmp_path / "test.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text(
        (72, 72),
        "Hello, World!",
        fontname="helv",
        fontsize=12,
    )
    page.insert_text(
        (72, 100),
        "Second line with a number 42.",
        fontname="helv",
        fontsize=12,
    )
    doc.save(str(p))
    doc.close()

    text = load_any(p)
    assert "Hello, World!" in text
    assert "42" in text


def test_load_docx_handles_tables(tmp_path: Path):
    """Make sure table cells get extracted too."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "خلية ١"
    table.cell(0, 1).text = "خلية ٢"
    table.cell(1, 0).text = "خلية ٣"
    p = tmp_path / "tab.docx"
    doc.save(str(p))

    text = load_any(p)
    assert "خلية ١" in text
    assert "خلية ٢" in text
