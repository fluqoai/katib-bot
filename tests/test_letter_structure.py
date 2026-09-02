from io import BytesIO
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor

from rag.export import build_official_branded_letter
from rag.generator import (
    FIXED_CLOSING,
    FIXED_FOOTER,
    GeneratedDraft,
    build_recipient_line,
    enforce_official_structure,
)


def _draft(body: str) -> GeneratedDraft:
    return GeneratedDraft(body=body, body_only=body, sources_block="")


def test_individual_opening():
    assert build_recipient_line({
        "entity_type": "individual",
        "recipient_name": "مدير فرع الوزارة",
    }) == "إلى سعادة مدير فرع الوزارة حفظه الله"


def test_organization_opening_does_not_duplicate_company_word():
    assert build_recipient_line({
        "entity_type": "organization",
        "recipient_name": "شركة تمكين التقنية",
    }) == "المكرمون شركة / تمكين التقنية سلمهم الله"


def test_structure_is_locked_and_deduplicated():
    result = enforce_official_structure(
        "الموضوع: طلب شراكة\n\nأما بعد،\n\nنرغب في التعاون.\n\nوتقبلوا وافر التحية والتقدير",
        known_fields={
            "entity_type": "organization",
            "recipient_name": "مؤسسة الخير",
        },
    )
    parts = result.split("\n\n")
    assert parts == [
        "المكرمون شركة / الخير سلمهم الله",
        "الموضوع: طلب شراكة",
        "السلام عليكم ورحمة الله وبركاته،",
        "أما بعد،",
        "نرغب في التعاون.",
        FIXED_CLOSING,
        FIXED_FOOTER,
    ]


def test_branded_export_preserves_header_and_footer_media():
    template = Path(__file__).resolve().parent.parent / "assets" / "official-letter-template.docx"
    body = enforce_official_structure(
        "الموضوع: خطاب تجريبي\nأما بعد،\nهذا متن تجريبي.",
        known_fields={"entity_type": "individual", "recipient_name": "مدير الفرع"},
    )
    data, _ = build_official_branded_letter(template, _draft(body))
    document = Document(BytesIO(data))
    text = [p.text for p in document.paragraphs if p.text.strip()]
    assert text[0] == "إلى سعادة مدير الفرع حفظه الله"
    assert text[-2:] == [FIXED_CLOSING, FIXED_FOOTER]
    visible_runs = [
        run
        for paragraph in document.paragraphs
        for run in paragraph.runs
        if run.text.strip()
    ]
    assert visible_runs
    assert all(run.font.color.rgb == RGBColor(0x33, 0x57, 0x6F) for run in visible_runs)
    assert all(run.font.color.theme_color is None for run in visible_runs)
    # Word mirrors physical left/right alignment for bidi paragraphs; LEFT in
    # the XML is the visual right edge in the rendered Arabic document.
    assert document.paragraphs[0].alignment == WD_PARAGRAPH_ALIGNMENT.LEFT
    assert document.paragraphs[1].alignment == WD_PARAGRAPH_ALIGNMENT.LEFT
    assert document.paragraphs[2].alignment == WD_PARAGRAPH_ALIGNMENT.LEFT
    assert document.paragraphs[3].alignment == WD_PARAGRAPH_ALIGNMENT.LEFT
    assert document.paragraphs[4].alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    assert document.paragraphs[-2].alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
    assert document.paragraphs[-1].alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
    assert all(run.font.name == "Times New Roman" for run in visible_runs)
    assert document.paragraphs[4].runs[0].font.size.pt == 15.0

    w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    assert all(p._p.pPr.find(f"{w_ns}bidi") is not None for p in document.paragraphs)
    assert all(run._r.rPr.find(f"{w_ns}rtl") is not None for run in visible_runs)

    with ZipFile(template) as source_package, ZipFile(BytesIO(data)) as package:
        anchors = []
        namespace = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
        for name in package.namelist():
            if name.startswith("word/header") and name.endswith(".xml"):
                anchors.extend(ElementTree.fromstring(package.read(name)).iter(f"{namespace}anchor"))
        for image_name in ("word/media/image1.png", "word/media/image2.png"):
            assert package.read(image_name) == source_package.read(image_name)
    assert len(anchors) == 2
    assert len(data) > 300_000  # branded images remain embedded
