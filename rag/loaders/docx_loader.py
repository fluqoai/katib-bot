"""DOCX loader.

Strategy:
  1. Try python-docx for paragraphs + tables (in document order).
  2. If that fails or yields empty text, the file is likely a renamed
     .doc (old binary format). Set status='needs_doc_conversion' so
     the caller can re-run via LibreOffice.
  3. Headers and footers are included because policy/legal docs sometimes
     carry the only metadata (e.g. "اعتماد مجلس الإدارة"). They're
     bracketed so the chunker can keep them associated with the body.
"""

from __future__ import annotations

import logging
from pathlib import Path

from rag.loaders.types import LoadedDoc


logger = logging.getLogger(__name__)


def load_docx(path: Path) -> LoadedDoc:
    """Read a .docx file: paragraphs, tables, headers, footers, in order.

    Returns a LoadedDoc with status='needs_doc_conversion' if the file
    is actually an old-format .doc renamed to .docx.
    """
    try:
        from docx import Document
    except ImportError:
        return LoadedDoc(
            text="",
            status="error",
            mime_type=None,
            error_message="python-docx is not installed",
            extractor="docx",
        )

    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    try:
        doc = Document(str(path))
    except Exception as e:  # noqa: BLE001
        # python-docx raises PackageNotFoundError for non-zip files (i.e.
        # old binary .doc files renamed to .docx). Convert to a clean
        # error that the dispatcher can handle.
        return LoadedDoc(
            text="",
            status="needs_doc_conversion",
            mime_type=mime,
            error_message=(
                f"python-docx cannot open the file ({type(e).__name__}). "
                "It looks like an old-format .doc, not a real .docx. "
                "Run the ingestion via LibreOffice to convert it first."
            ),
            extractor="docx",
        )

    parts: list[str] = []

    # Headers/footers — only ONCE per document, deduped. Some DOCX
    # files (especially policy templates with many sections) repeat
    # the same header in every section, which would balloon the text
    # and explode the chunk count. We collect unique paragraphs.
    seen_header: set[str] = set()
    seen_footer: set[str] = set()
    for section in doc.sections:
        for hdr in (section.header, section.first_page_header, section.even_page_header):
            if hdr is None:
                continue
            for p in hdr.paragraphs:
                t = p.text.strip()
                if t and t not in seen_header:
                    seen_header.add(t)
                    parts.append(f"[HEADER] {t}")
        for ftr in (section.footer, section.first_page_footer, section.even_page_footer):
            if ftr is None:
                continue
            for p in ftr.paragraphs:
                t = p.text.strip()
                if t and t not in seen_footer:
                    seen_footer.add(t)
                    parts.append(f"[FOOTER] {t}")

    # Body — iterate the body in document order so we get paragraphs and
    # tables interleaved correctly (python-docx exposes them as separate
    # collections; we use `iter_inner_content` to interleave).
    try:
        for child in doc.element.body.iterchildren():
            tag = child.tag.rsplit("}", 1)[-1]
            if tag == "p":
                # Find the matching paragraph object
                for p in doc.paragraphs:
                    if p._element is child:  # type: ignore[attr-defined]
                        t = p.text.strip()
                        if t:
                            parts.append(t)
                        break
            elif tag == "tbl":
                for table in doc.tables:
                    if table._element is child:  # type: ignore[attr-defined]
                        for row in table.rows:
                            row_text = " | ".join(
                                cell.text.strip() for cell in row.cells
                            )
                            if row_text.strip(" |"):
                                parts.append(row_text)
                        break
    except Exception as e:  # noqa: BLE001
        logger.warning("Body iteration failed for %s: %s", path, e)
        # Fallback: just paragraphs + tables in order
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    parts.append(row_text)

    text = "\n\n".join(parts).strip()
    if not text:
        return LoadedDoc(
            text="",
            status="empty",
            mime_type=mime,
            error_message="file opened but contains no extractable text",
            extractor="docx",
        )

    return LoadedDoc(
        text=text,
        status="ok",
        mime_type=mime,
        extractor="docx",
    )
