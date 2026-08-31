"""DOCX + PDF export for the letter pipeline.

DOCX generation treats the uploaded template as a *style source*,
not as content to keep alongside the new letter. The output is a
fresh, official-looking letter that adopts the template's:

  - RTL direction
  - fonts (default + theme)
  - page setup (margins, size, orientation)
  - sections (headers, footers, page numbers)
  - styles (paragraph, character, table)
  - tables, lists, and other structural elements

Three template modes are supported, in priority order:

  1. ``{{body}}`` marker present
     The generated letter is substituted into that exact location.
     All other paragraphs of the template are kept untouched (logo,
     header, signature block, footer, etc.).
  2. Placeholder fields present (``{{recipient_name}}`` etc.)
     The placeholders are filled in place. The body is written after
     the last existing paragraph (still using the template's styles).
  3. No markers (template is a complete example)
     The template's body is **cleared** and the new letter is written
     from scratch using its styles. The original letter's text is
     discarded; only the formatting is kept.

The final DOCX is an official, ready-to-send document. It must NOT
contain the ``[source: ...]`` inline tags or the ``## المصادر``
section — those are for the internal provenance record only.

PDF is produced by headless LibreOffice. If LibreOffice isn't
installed, the export returns ``pdf_available=False`` and only the
DOCX is produced. The caller decides what to do (return null URLs,
surface an error, etc.).
"""
from __future__ import annotations

import io
import logging
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

from rag.evidence import EvidenceBundle, SourceChunk
from rag.generator import GeneratedDraft


logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------

# Marker recognised as "insert the full body here". Templates that
# have this marker keep all their surrounding content (logo, header,
# signature block, etc.) and only the body text is swapped.
BODY_MARKER = "{{body}}"


@dataclass(slots=True)
class TemplateProfile:
    """Snapshot of the visual style of a template DOCX.

    Captured by :func:`extract_template_profile`. We apply this to a
    fresh document so the output letter inherits the template's look
    without inheriting its text.
    """
    template_path: Path
    has_body_marker: bool
    placeholder_names: list[str]            # e.g. ['recipient_name', 'partner_org']
    rtl: bool                              # True for Arabic/Hebrew templates
    paragraph_count: int
    section_count: int


def extract_template_profile(template_path: Path) -> TemplateProfile:
    """Open a DOCX and return its visual profile + structural hints."""
    try:
        from docx import Document
    except ImportError as e:  # noqa: BLE001
        raise RuntimeError("python-docx is required for DOCX export") from e

    doc = Document(str(template_path))

    has_body_marker = any(BODY_MARKER in p.text for p in doc.paragraphs) or any(
        BODY_MARKER in cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )

    # Collect every `{{name}}` placeholder that appears in the template.
    placeholder_pat = re.compile(r"\{\{\s*([a-zA-Z0-9_\-\u0600-\u06FF]+)\s*\}\}")
    found: set[str] = set()
    for p in doc.paragraphs:
        found.update(placeholder_pat.findall(p.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    found.update(placeholder_pat.findall(p.text))

    # RTL: any paragraph with right-to-left alignment or any text that
    # contains Arabic characters. Arabic templates are always RTL.
    rtl = any(
        p.alignment is not None and int(p.alignment) == 2  # WD_ALIGN_PARAGRAPH.RIGHT
        for p in doc.paragraphs
    )
    if not rtl:
        for p in doc.paragraphs:
            if re.search(r"[\u0600-\u06FF]", p.text):
                rtl = True
                break

    return TemplateProfile(
        template_path=template_path,
        has_body_marker=has_body_marker,
        placeholder_names=sorted(found),
        rtl=rtl,
        paragraph_count=len(doc.paragraphs),
        section_count=len(doc.sections),
    )


def _copy_template_sections(src_doc, dst_doc) -> None:
    """Copy page setup (margins, size, orientation) from src to dst sections."""
    from docx.shared import Pt
    from docx.enum.section import WD_ORIENTATION

    # python-docx creates a default section on construction; we want
    # to match the count of the source.
    src_sections = list(src_doc.sections)
    # Ensure dst has at least one section (it always does)
    while len(dst_doc.sections) < len(src_sections):
        dst_doc.add_section()
    # Copy each section's page setup
    for i, src in enumerate(src_sections):
        if i >= len(dst_doc.sections):
            break
        dst = dst_doc.sections[i]
        dst.page_height = src.page_height
        dst.page_width  = src.page_width
        dst.top_margin    = src.top_margin
        dst.bottom_margin = src.bottom_margin
        dst.left_margin   = src.left_margin
        dst.right_margin  = src.right_margin
        dst.header_distance = src.header_distance
        dst.footer_distance = src.footer_distance
        dst.orientation  = src.orientation
        # Headers / footers
        try:
            dst.header.is_linked_to_previous = False
        except Exception:  # noqa: BLE001
            pass
        try:
            dst.footer.is_linked_to_previous = False
        except Exception:  # noqa: BLE001
            pass


def _strip_citation_tags(text: str) -> str:
    """Remove ``[source: label]`` tags from the body before writing.

    The DOCX is an official document and must not show internal
    citation markers. Sources are persisted separately in the
    `drafts.sources` jsonb column.
    """
    return re.sub(r"\s*\[source:\s*[^\]]+\]", "", text).strip()


# OOXML namespace for w: tags (used to set bidi / rtl directly)
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _apply_rtl_to_paragraph(paragraph, *, font_size_pt: float | None = None,
                            bold: bool = False, space_after_pt: float | None = 6.0) -> None:
    """Apply explicit Arabic RTL formatting to a paragraph and its runs.

    This sets the OOXML properties that are required for unambiguous
    RTL rendering in any Word-compatible viewer (not relying on
    Unicode bidi heuristics alone):

      - <w:pPr><w:bidi/></w:pPr>          paragraph direction RTL
      - <w:rPr><w:rtl/></w:rPr>          Unicode RTL on each run
      - alignment: right                 visual right-to-left alignment
      - font size (optional)             preserve template size
      - bold (optional)                  for headings / labels
      - <w:spacing w:after="..."/>       paragraph spacing

    The Word UI lets users override these per-document, so the file
    remains editable in any language.
    """
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from lxml import etree

    pf = paragraph.paragraph_format
    pf.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    if space_after_pt is not None:
        pf.space_after = Pt(space_after_pt)

    # Set <w:bidi/> in <w:pPr> for explicit paragraph direction
    pPr = paragraph._p.get_or_add_pPr()
    # Remove any existing bidi tag
    for existing in pPr.findall(f"{{{W_NS}}}bidi"):
        pPr.remove(existing)
    bidi = etree.SubElement(pPr, f"{{{W_NS}}}bidi")
    bidi.set(f"{{{W_NS}}}val", "1")

    # Apply font / bold / RTL to every run
    for run in paragraph.runs:
        if font_size_pt is not None:
            run.font.size = Pt(font_size_pt)
        if bold:
            run.bold = True
        # Set <w:rtl/> in run's rPr
        rPr = run._r.get_or_add_rPr()
        for existing in rPr.findall(f"{{{W_NS}}}rtl"):
            rPr.remove(existing)
        rtl = etree.SubElement(rPr, f"{{{W_NS}}}rtl")
        rtl.set(f"{{{W_NS}}}val", "1")


def _strip_markdown_to_paragraphs(text: str) -> list[tuple[str, str]]:
    """Convert a lightly-marked-up body into (style, text) tuples.

    Recognised markers:
      * Lines starting with ``##`` → heading level 2
      * Lines starting with ``#``  → heading level 1
      * Lines starting with ``-`` or ``*`` → list item
      * Lines that look like ``**...**`` → bold
      * Empty lines → blank paragraph (spacing)
    Returns a list suitable for ``doc.add_paragraph``.
    """
    out: list[tuple[str, str]] = []
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line.strip():
            out.append(("normal", ""))
            continue
        if line.startswith("## "):
            out.append(("h2", line[3:].strip()))
        elif line.startswith("# "):
            out.append(("h1", line[2:].strip()))
        elif line.startswith(("- ", "* ", "• ")):
            out.append(("list", line[2:].strip()))
        else:
            out.append(("normal", line))
    return out


def _find_arabic_body_block(paragraphs) -> tuple[int, int]:
    """Find the body block in a "complete example" template.

    In a formal Arabic letter, the paragraphs that should be REPLACED
    (the actual body content) are the long, narrative paragraphs in
    the MIDDLE. The other Arabic-text paragraphs are typically:

    * Recipient line         — short, LEFT-aligned
    * Greeting ("السلام عليكم") — short, CENTER-aligned
    * Closing ("شاكرين ومقدرين") — CENTER-aligned
    * Farewell ("والله يحفظكم")  — short, CENTER-aligned
    * Signature              — usually CENTER or just a name

    Heuristic:
      1. Find paragraphs with substantial Arabic text (≥ 30 Arabic chars)
      2. Exclude paragraphs whose explicit alignment is CENTER or LEFT
         (those are header/recipient/greeting/closing/signature slots)
      3. Among the remaining, find the longest contiguous run
      4. Fall back to the simple "longest Arabic run" if step 3 yields
         nothing (template with all-CENTER body, etc.)

    Returns ``(start_index, end_index_exclusive)``. If no Arabic
    paragraphs are found, returns ``(0, 0)``.
    """
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

    MIN_BODY_CHARS = 30
    CENTER = WD_PARAGRAPH_ALIGNMENT.CENTER
    LEFT = WD_PARAGRAPH_ALIGNMENT.LEFT

    def is_body_candidate(p) -> bool:
        text = p.text or ""
        arabic_chars = sum(1 for c in text if '\u0600' <= c <= '\u06FF')
        if arabic_chars < MIN_BODY_CHARS:
            return False
        align = p.alignment
        if align in (CENTER, LEFT):
            return False
        return True

    # 1) alignment-aware run
    body_idx = [i for i, p in enumerate(paragraphs) if is_body_candidate(p)]
    if body_idx:
        best_start, best_len = body_idx[0], 1
        cur_start, cur_len = body_idx[0], 1
        for k in range(1, len(body_idx)):
            if body_idx[k] == body_idx[k - 1] + 1:
                cur_len += 1
            else:
                if cur_len > best_len:
                    best_start, best_len = cur_start, cur_len
                cur_start = body_idx[k]
                cur_len = 1
        if cur_len > best_len:
            best_start, best_len = cur_start, cur_len
        if best_len > 0:
            return (best_start, best_start + best_len)

    # 2) fallback: longest Arabic run (ignores alignment)
    arabic_idx = [i for i, p in enumerate(paragraphs)
                  if re.search(r"[\u0600-\u06FF]", p.text)]
    if not arabic_idx:
        return (0, 0)
    best_start, best_len = arabic_idx[0], 1
    cur_start, cur_len = arabic_idx[0], 1
    for k in range(1, len(arabic_idx)):
        if arabic_idx[k] == arabic_idx[k - 1] + 1:
            cur_len += 1
        else:
            if cur_len > best_len:
                best_start, best_len = cur_start, cur_len
            cur_start = arabic_idx[k]
            cur_len = 1
    if cur_len > best_len:
        best_start, best_len = cur_start, cur_len
    return (best_start, best_start + best_len)


def _clone_paragraph_after(doc, anchor_para):
    """Create a new paragraph immediately after ``anchor_para`` that
    copies its ``pPr`` (alignment, spacing, indent, bidi, rPr-on-pPr,
    etc.) but starts with no runs. The caller sets the text.

    Used when the new body is longer than the template's body block
    and we need extra paragraphs that visually match the template.
    """
    from copy import deepcopy
    from docx.text.paragraph import Paragraph
    new_p_el = deepcopy(anchor_para._p)
    # Strip runs + hyperlinks; keep pPr so the clone inherits all
    # paragraph-level formatting (alignment, spacing, indent, bidi).
    for tag in (f"{{{W_NS}}}r", f"{{{W_NS}}}hyperlink"):
        for el in new_p_el.findall(tag):
            new_p_el.remove(el)
    anchor_para._p.addnext(new_p_el)
    return Paragraph(new_p_el, anchor_para._parent)


def _set_paragraph_text(paragraph, text: str) -> None:
    """Set a paragraph's text while keeping its ``pPr`` (alignment,
    spacing, indent, bidi, line-spacing) and the style-defined default
    run font. We do this by replacing the runs directly rather than
    using python-docx's ``paragraph.text =`` setter, so any pPr-only
    XML (no rPr at the paragraph level) is preserved verbatim.

    If the paragraph's ``pPr`` does not already have ``<w:bidi/>`` and
    the new text contains Arabic, we add an explicit ``<w:bidi/>`` so
    the paragraph renders unambiguously RTL in every Word viewer
    (and not just via Unicode bidi heuristics). Same for the new run:
    if it has Arabic and no existing ``<w:rtl/>``, we add one.
    """
    from lxml import etree
    p_el = paragraph._p
    # Remove existing runs + hyperlinks; keep pPr.
    for tag in (f"{{{W_NS}}}r", f"{{{W_NS}}}hyperlink"):
        for el in p_el.findall(tag):
            p_el.remove(el)

    text_has_arabic = bool(re.search(r"[\u0600-\u06FF]", text or ""))

    # Add <w:bidi/> to pPr if it's missing and the text is Arabic.
    if text_has_arabic:
        pPr = p_el.find(f"{{{W_NS}}}pPr")
        if pPr is not None and pPr.find(f"{{{W_NS}}}bidi") is None:
            etree.SubElement(pPr, f"{{{W_NS}}}bidi")

    # Add a single new run with the text. If Arabic and no rPr yet,
    # add an rPr with <w:rtl/> so the run is unambiguously RTL.
    r = etree.SubElement(p_el, f"{{{W_NS}}}r")
    if text_has_arabic:
        rPr = etree.SubElement(r, f"{{{W_NS}}}rPr")
        etree.SubElement(rPr, f"{{{W_NS}}}rtl")
    t = etree.SubElement(r, f"{{{W_NS}}}t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text


def build_letter_from_template(
    template_path: Path,
    draft: GeneratedDraft,
    *,
    placeholders: dict[str, str] | None = None,
    include_sources: bool = False,           # default False: official letter stays clean
    out_path: Path | None = None,
) -> tuple[bytes, TemplateProfile]:
    """Build the final DOCX from a template + the generated draft.

    The template is treated as a *style source* — see the module
    docstring for the three modes (body marker, placeholder fields,
    complete-example). The returned bytes are an OFFICIAL, ready-to-
    send letter; the body has its ``[source: ...]`` tags removed and
    the ``## المصادر`` section is omitted unless ``include_sources``
    is True (intended for internal review only).

    Returns ``(docx_bytes, profile)`` so the caller can report which
    mode was used and which placeholders were filled.
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError as e:  # noqa: BLE001
        raise RuntimeError("python-docx is required for DOCX export") from e

    placeholders = placeholders or {}
    profile = extract_template_profile(template_path)
    src = Document(str(template_path))
    dst = Document(str(template_path))  # start from a copy of the template

    # Clean body text of the draft
    body_text = _strip_citation_tags(draft.body_only)

    # ------------------------------------------------------------------
    # Mode 1: {{body}} marker → swap in place, keep everything else
    # ------------------------------------------------------------------
    if profile.has_body_marker:
        for p in dst.paragraphs:
            if BODY_MARKER in p.text:
                _replace_in_runs(p, BODY_MARKER, body_text)
                break
    else:
        # ------------------------------------------------------------------
        # Mode 2 & 3: fill placeholders, then either keep the body or
        # clear it and rewrite using template styles.
        # ------------------------------------------------------------------
        # Fill {{placeholder}} tokens everywhere they appear
        if profile.placeholder_names:
            for k in profile.placeholder_names:
                token = "{{" + k + "}}"
                value = placeholders.get(k, "") or f"({k})"
                for p in dst.paragraphs:
                    if token in p.text:
                        _replace_in_runs(p, token, value)
                for table in dst.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if token in p.text:
                                    _replace_in_runs(p, token, value)

        # Decide: keep the template's body as-is, or rewrite it.
        # If the template is clearly a "complete example" (no placeholders
        # AND no body marker, AND has substantial body text), we treat
        # it as Mode 3: clear body and rewrite using template styles.
        is_complete_example = (
            not profile.placeholder_names
            and any(re.search(r"[\u0600-\u06FF]", p.text) for p in dst.paragraphs)
        )
        if is_complete_example:
            # IN-PLACE body replacement: keep the template's structure
            # (top spacing, basmala, recipient, greeting, closing,
            # signature, tables, images, custom paragraph styles) and
            # only swap the text inside the body block. The body
            # block is the longest contiguous run of Arabic-text
            # paragraphs in the template.
            body_start, body_end = _find_arabic_body_block(dst.paragraphs)
            body_block = list(dst.paragraphs[body_start:body_end])

            if not body_block:
                # Defensive fallback: Arabic text was reported earlier
                # (is_complete_example=True) but the body block ended
                # up empty for some reason. In that case we still need
                # to produce output, so fall back to the original
                # wipe+rewrite behaviour.
                body = dst.element.body
                sectPr = body.find(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr"
                )
                for child in list(body):
                    if child.tag.endswith("}sectPr"):
                        continue
                    body.remove(child)
                for style, text in _strip_markdown_to_paragraphs(body_text):
                    if not text:
                        p = dst.add_paragraph("")
                        _apply_rtl_to_paragraph(p, space_after_pt=3.0)
                        continue
                    if style == "h1":
                        p = _add_styled_heading(dst, text, level=1)
                        _apply_rtl_to_paragraph(p, font_size_pt=16, bold=True, space_after_pt=8.0)
                    elif style == "h2":
                        p = _add_styled_heading(dst, text, level=2)
                        _apply_rtl_to_paragraph(p, font_size_pt=14, bold=True, space_after_pt=6.0)
                    elif style == "list":
                        p = _safe_add_paragraph(dst, text, style_name="List Bullet")
                        _apply_rtl_to_paragraph(p, space_after_pt=3.0)
                    else:
                        p = dst.add_paragraph(text)
                        _apply_rtl_to_paragraph(p, font_size_pt=12, space_after_pt=6.0)
            else:
                # Parse the new body into plain lines. We do not apply
                # markdown styling here — the template's existing
                # paragraph properties (alignment, font, size, spacing,
                # bidi) are what we want the new content to inherit.
                # A formal Arabic letter is a sequence of plain
                # paragraphs, not a markdown document.
                new_lines: list[str] = []
                for raw in body_text.splitlines():
                    s = raw.strip()
                    if not s:
                        continue
                    if s.startswith("## "):
                        s = s[3:]
                    elif s.startswith("# "):
                        s = s[2:]
                    elif s.startswith(("- ", "* ", "• ")):
                        s = s[2:]
                    new_lines.append(s)

                # Replace text in the template's body paragraphs, one
                # line per paragraph. Paragraphs outside [body_start,
                # body_end) are NOT touched.
                for i, p in enumerate(body_block):
                    if i < len(new_lines):
                        _set_paragraph_text(p, new_lines[i])
                    else:
                        _set_paragraph_text(p, "")

                # If the new body is longer than the body block, clone
                # the last body paragraph's pPr to make new paragraphs
                # that visually match the template. Insert them in
                # order, immediately after the last body paragraph, so
                # they stay inside the body block.
                if len(new_lines) > len(body_block):
                    last_para = body_block[-1]
                    for i in range(len(body_block), len(new_lines)):
                        new_p = _clone_paragraph_after(dst, last_para)
                        _set_paragraph_text(new_p, new_lines[i])
                        last_para = new_p
                        body_block.append(new_p)
        else:
            # Mode 2: placeholders filled, but template is sparse. Append
            # the new body at the end of the existing content using the
            # template's default styles. Same RTL treatment as Mode 3.
            if dst.paragraphs:
                # add a separator
                sep = dst.add_paragraph("")
                _apply_rtl_to_paragraph(sep, space_after_pt=6.0)
            for style, text in _strip_markdown_to_paragraphs(body_text):
                if not text:
                    p = dst.add_paragraph("")
                    _apply_rtl_to_paragraph(p, space_after_pt=3.0)
                    continue
                if style == "h1":
                    p = _add_styled_heading(dst, text, level=1)
                    _apply_rtl_to_paragraph(p, font_size_pt=16, bold=True, space_after_pt=8.0)
                elif style == "h2":
                    p = _add_styled_heading(dst, text, level=2)
                    _apply_rtl_to_paragraph(p, font_size_pt=14, bold=True, space_after_pt=6.0)
                elif style == "list":
                    p = _safe_add_paragraph(dst, text, style_name="List Bullet")
                    _apply_rtl_to_paragraph(p, space_after_pt=3.0)
                else:
                    p = dst.add_paragraph(text)
                    _apply_rtl_to_paragraph(p, font_size_pt=12, space_after_pt=6.0)

    # Optionally append the sources section (for internal review only)
    if include_sources and (draft.sources_block or draft.parsed_citations):
        dst.add_paragraph("")
        dst.add_paragraph("## المصادر")
        for cit in draft.parsed_citations:
            dst.add_paragraph(
                f"[{cit.get('label', '?')}]: {cit.get('title', '?')} — chunk "
                f"#{cit.get('chunk_index', '?')} — similarity={cit.get('similarity', '?')}"
            )

    # Save
    if out_path is None:
        buf = io.BytesIO()
        dst.save(buf)
        return buf.getvalue(), profile
    dst.save(str(out_path))
    return out_path.read_bytes(), profile


def _doc_has_text(doc, needle: str) -> bool:
    return any(needle in p.text for p in doc.paragraphs) or any(
        needle in cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )


def _add_styled_heading(doc, text: str, *, level: int = 2) -> None:
    """Add a heading paragraph using direct formatting (bold + larger
    font). This is style-name-independent — works even when the
    template's styles.xml doesn't define a "Heading 2" style.

    Heading 1: bold, 16pt, with bottom spacing
    Heading 2: bold, 14pt, with bottom spacing
    """
    from docx.shared import Pt
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16 if level <= 1 else 14)
    # Set a small space after the heading
    try:
        p.paragraph_format.space_after = Pt(6)
    except Exception:  # noqa: BLE001
        pass
    return p


def _safe_add_paragraph(doc, text: str, *, style_name: str | None = None):
    """Add a paragraph, falling back to default style if the named
    style isn't defined. ``add_paragraph(text, style="List Bullet")``
    raises KeyError when the template has no such style; we don't
    want a missing style to abort the whole export.
    """
    if not style_name:
        return doc.add_paragraph(text)
    try:
        return doc.add_paragraph(text, style=style_name)
    except KeyError:
        logger.warning("style %r not in template; using default", style_name)
        return doc.add_paragraph(text)


def _replace_in_runs(paragraph, old: str, new: str) -> None:
    """Replace ``old`` with ``new`` in a paragraph, handling the
    run-splitting issue. python-docx splits text across multiple
    runs; a naive ``.text =`` would destroy formatting. We rebuild
    the paragraph keeping the first run's formatting.
    """
    if not paragraph.runs:
        paragraph.add_run(new)
        return
    full = "".join(r.text for r in paragraph.runs)
    if old not in full:
        return
    new_full = full.replace(old, new)
    first = paragraph.runs[0]
    for r in paragraph.runs[1:]:
        r.text = ""
    first.text = new_full


# Backwards-compatible alias for the old single-mode function.
def build_docx(
    template_path: Path,
    draft: GeneratedDraft,
    *,
    placeholders: dict[str, str] | None = None,
    out_path: Path | None = None,
) -> bytes:
    """Backwards-compatible wrapper.

    Kept for callers that don't care about the new modes. Always
    produces a clean letter (no inline citation tags, no
    ``## المصادر`` section) — that is the desired default.
    """
    data, _profile = build_letter_from_template(
        template_path,
        draft,
        placeholders=placeholders,
        include_sources=False,
        out_path=out_path,
    )
    return data


# ---------------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------------

def has_soffice() -> bool:
    """True iff LibreOffice's ``soffice`` binary is available on this host."""
    if shutil.which("soffice") or shutil.which("libreoffice"):
        return True
    for p in (
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ):
        if Path(p).exists():
            return True
    return False


def build_pdf(docx_bytes: bytes, out_path: Path | None = None) -> bytes | None:
    """Convert DOCX bytes to PDF using headless LibreOffice.

    Returns the PDF bytes on success. Returns ``None`` if LibreOffice
    is not installed or the conversion fails — the caller should
    surface ``pdf_available=False`` to the user, **never** silently
    return the DOCX bytes as a PDF.
    """
    if not has_soffice():
        logger.warning("soffice not found — PDF conversion unavailable")
        return None

    soffice = (
        shutil.which("soffice")
        or shutil.which("libreoffice")
        or r"C:\Program Files\LibreOffice\program\soffice.exe"
    )

    with tempfile.TemporaryDirectory(prefix="kateb-pdf-") as tmp:
        tmp_path = Path(tmp)
        docx_file = tmp_path / "input.docx"
        docx_file.write_bytes(docx_bytes)
        try:
            proc = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf",
                 "--outdir", str(tmp_path), str(docx_file)],
                check=False, capture_output=True, text=True, timeout=120,
            )
        except subprocess.TimeoutExpired:
            logger.error("soffice convert timed out")
            return None
        if proc.returncode != 0:
            logger.error("soffice convert failed: %s", proc.stderr or proc.stdout)
            return None
        pdf_path = tmp_path / "input.pdf"
        if not pdf_path.exists():
            logger.error("soffice did not produce a PDF in %s", tmp_path)
            return None
        data = pdf_path.read_bytes()
        if out_path:
            out_path.write_bytes(data)
        return data


__all__ = [
    "build_docx",
    "build_letter_from_template",
    "build_pdf",
    "extract_template_profile",
    "has_soffice",
    "TemplateProfile",
    # --- Phase 1: regulation-driven, template-independent export (NEW) ---
    "LetterStyle",
    "DEFAULT_LETTER_STYLE",
    "apply_letter_style",
    "build_letter_from_style",
    "build_letter",
]


# ---------------------------------------------------------------------------
# Phase 1 — Template-independent DOCX export
#
# This block adds a new way to produce a DOCX letter that does NOT
# require a template file. The visual style is described by a
# ``LetterStyle`` dataclass; the function ``build_letter_from_style``
# constructs a fresh DOCX from a blank document and applies the style.
#
# Why: the pipeline now treats templates as OPTIONAL hints. When no
# template is available (or the user opts out of legacy templates), the
# export must still produce a clean, official-looking letter.
#
# The legacy path (``build_letter_from_template``) is untouched and
# remains the default for any caller that passes a template path.
# ---------------------------------------------------------------------------


@dataclass(slots=True)
class LetterStyle:
    """Visual style for an official Saudi administrative letter,
    independent of any template DOCX file.

    The pipeline builds a fresh DOCX from this spec + the LLM-generated
    body. No template file is required.

    All defaults are conservative — they produce a clean, formal
    Arabic letter suitable for any association context.
    """
    # Page
    page_size: str = "A4"                  # "A4" | "Letter"
    margin_top_cm: float = 2.5
    margin_bottom_cm: float = 2.5
    margin_left_cm: float = 3.0
    margin_right_cm: float = 3.0

    # Direction
    rtl: bool = True                       # Arabic: always True in practice

    # Typography
    font_name: str = "Arial"               # Latin font; Arabic glyphs are rendered by the OS
    font_name_arabic: str = "Traditional Arabic"  # hint for the Arabic glyphs
    font_size_pt: float = 14.0             # default body size
    line_spacing: float = 1.15

    # Paragraph spacing
    paragraph_spacing_pt: float = 6.0

    # Letter skeleton (which sections to include)
    include_basmala: bool = True           # "بسم الله الرحمن الرحيم" centered
    include_date_row: bool = True          # التاريخ / الموافق placeholder lines
    include_recipient_block: bool = True   # "سعادة / (...) المحترم"
    include_signature_block: bool = True   # signature + title lines
    include_closing_phrase: bool = True    # "وتقبلوا وافر التحية والتقدير،،،"

    # Closing phrase text
    closing_phrase: str = "وتقبلوا وافر التحية والتقدير،،،"
    signature_title: str = "رئيس مجلس الإدارة"
    signature_name_placeholder: str = "(يُستكمل اسم المُوقِّع)"

    # Date row text (placeholders that the LLM/generator may fill in
    # if known_fields include the values, or remain for the user to fill)
    date_label: str = "التاريخ"
    date_gregorian_placeholder: str = "..../..../......."
    date_hijri_label: str = "الموافق"
    date_hijri_placeholder: str = "..../..../......."

    # Opening
    recipient_prefix: str = "سعادة /"
    recipient_placeholder: str = "(يُستكمل اسم المُراسَل)"
    recipient_suffix: str = "المحترم،"

    # Basmala text
    basmala_text: str = "بسم الله الرحمن الرحيم"

    # Greeting
    greeting_text: str = "السلام عليكم ورحمة الله وبركاته، وبعد:"

    # Header alignment for the body (Arabic reads right-to-left)
    body_alignment: str = "right"          # "right" | "center" | "left" | "justified"

    # Numerals
    use_arabic_numerals: bool = True


# Sensible default — used when an intent has no specific style entry.
DEFAULT_LETTER_STYLE = LetterStyle()


def apply_letter_style(doc, style: LetterStyle) -> None:
    """Apply a ``LetterStyle`` to a blank (or fresh) ``Document``.

    Sets:
      * page size + margins
      * default font (Latin name + Arabic hint) on the Normal style
      * default paragraph spacing + line spacing on the Normal style
      * adds an explicit ``<w:bidi/>`` to the Normal pPr so any new
        paragraph without explicit formatting is still RTL-correct

    Per-paragraph bidi/rtl/alignment is applied by ``_apply_rtl_to_paragraph``
    which is reused here. This function only sets document-level defaults;
    the caller is responsible for applying per-paragraph RTL via the
    existing helper.
    """
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from lxml import etree

    # Page size
    if style.page_size.upper() == "A4":
        for section in doc.sections:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)
    elif style.page_size.upper() == "LETTER":
        for section in doc.sections:
            section.page_width = Cm(21.59)
            section.page_height = Cm(27.94)

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(style.margin_top_cm)
        section.bottom_margin = Cm(style.margin_bottom_cm)
        section.left_margin = Cm(style.margin_left_cm)
        section.right_margin = Cm(style.margin_right_cm)

    # Default font on the Normal style (Latin)
    normal = doc.styles["Normal"]
    normal.font.name = style.font_name
    normal.font.size = Pt(style.font_size_pt)

    # Default paragraph format (spacing + line spacing)
    normal.paragraph_format.space_after = Pt(style.paragraph_spacing_pt)
    normal.paragraph_format.line_spacing = style.line_spacing

    # Explicit <w:bidi/> on the Normal style's pPr so any new paragraph
    # created via doc.add_paragraph() inherits the right direction
    style_el = normal.element
    pPr = style_el.find(f"{{{W_NS}}}pPr")
    if pPr is None:
        pPr = etree.SubElement(style_el, f"{{{W_NS}}}pPr")
    for existing in pPr.findall(f"{{{W_NS}}}bidi"):
        pPr.remove(existing)
    bidi = etree.SubElement(pPr, f"{{{W_NS}}}bidi")
    bidi.set(f"{{{W_NS}}}val", "1")


def _build_letter_skeleton(doc, style: LetterStyle, draft: GeneratedDraft) -> None:
    """Write the standard Saudi administrative letter skeleton
    around the LLM-generated body.

    Layout (when all flags are True):
      [Date row]                  "التاريخ ....  الموافق ...."  (right-aligned)
      [Empty paragraph]           visual spacing
      [Basmala]                   "بسم الله الرحمن الرحيم"     (centered)
      [Empty paragraph]           visual spacing
      [Recipient block]           "سعادة / (...) المحترم،"     (right-aligned)
      [Greeting]                  "السلام عليكم ورحمة الله وبركاته، وبعد:"  (right-aligned)
      [Empty paragraph]           visual spacing
      [BODY]                      from the LLM draft (right-aligned, body paragraphs)
      [Empty paragraph]           visual spacing
      [Closing phrase]            "وتقبلوا وافر التحية والتقدير،،،"  (right-aligned)
      [Empty paragraph]           visual spacing
      [Signature block]           "(يُستكمل اسم المُوقِّع)" / "رئيس مجلس الإدارة"  (centered)

    Each paragraph receives an explicit ``<w:bidi/>`` + ``<w:rtl/>``
    via the existing ``_apply_rtl_to_paragraph`` helper.
    """
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from lxml import etree

    body_text = _strip_citation_tags(draft.body_only or draft.body or "")

    # Date row
    if style.include_date_row:
        date_line = (
            f"{style.date_label}: {style.date_gregorian_placeholder}    "
            f"{style.date_hijri_label}: {style.date_hijri_placeholder}"
        )
        p = doc.add_paragraph(date_line)
        _apply_rtl_to_paragraph(p, font_size_pt=style.font_size_pt,
                                space_after_pt=style.paragraph_spacing_pt)
        doc.add_paragraph("")  # spacing

    # Basmala (centered)
    if style.include_basmala:
        p = doc.add_paragraph(style.basmala_text)
        # Centered, bold, larger
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.paragraph_format.space_after = Pt(style.paragraph_spacing_pt)
        # Inject runs so the helper can set rtl on them
        if not p.runs:
            p.add_run(style.basmala_text)
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(style.font_size_pt + 2)
        # Apply bidi to the paragraph (alignment overrides RIGHT in helper)
        pPr = p._p.get_or_add_pPr()
        for existing in pPr.findall(f"{{{W_NS}}}bidi"):
            pPr.remove(existing)
        bidi = etree.SubElement(pPr, f"{{{W_NS}}}bidi")
        bidi.set(f"{{{W_NS}}}val", "1")
        for run in p.runs:
            rPr = run._r.get_or_add_rPr()
            for existing in rPr.findall(f"{{{W_NS}}}rtl"):
                rPr.remove(existing)
            rtl = etree.SubElement(rPr, f"{{{W_NS}}}rtl")
            rtl.set(f"{{{W_NS}}}val", "1")
        doc.add_paragraph("")  # spacing

    # Recipient block
    if style.include_recipient_block:
        recipient_line = (
            f"{style.recipient_prefix} {style.recipient_placeholder} {style.recipient_suffix}"
        )
        p = doc.add_paragraph(recipient_line)
        _apply_rtl_to_paragraph(p, font_size_pt=style.font_size_pt,
                                bold=True,
                                space_after_pt=style.paragraph_spacing_pt)
        # Greeting
        p = doc.add_paragraph(style.greeting_text)
        _apply_rtl_to_paragraph(p, font_size_pt=style.font_size_pt,
                                space_after_pt=style.paragraph_spacing_pt)
        doc.add_paragraph("")  # spacing

    # Body — strip markdown into paragraphs
    for style_name, text in _strip_markdown_to_paragraphs(body_text):
        if not text:
            # blank paragraph
            p = doc.add_paragraph("")
            _apply_rtl_to_paragraph(p, font_size_pt=style.font_size_pt,
                                    space_after_pt=style.paragraph_spacing_pt)
            continue
        if style_name == "h1":
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(style.font_size_pt + 4)
            # bidi + rtl
            pPr = p._p.get_or_add_pPr()
            for existing in pPr.findall(f"{{{W_NS}}}bidi"):
                pPr.remove(existing)
            bidi = etree.SubElement(pPr, f"{{{W_NS}}}bidi")
            bidi.set(f"{{{W_NS}}}val", "1")
            rPr = run._r.get_or_add_rPr()
            for existing in rPr.findall(f"{{{W_NS}}}rtl"):
                rPr.remove(existing)
            rtl = etree.SubElement(rPr, f"{{{W_NS}}}rtl")
            rtl.set(f"{{{W_NS}}}val", "1")
            p.paragraph_format.space_after = Pt(style.paragraph_spacing_pt)
        elif style_name == "h2":
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(style.font_size_pt + 2)
            pPr = p._p.get_or_add_pPr()
            for existing in pPr.findall(f"{{{W_NS}}}bidi"):
                pPr.remove(existing)
            bidi = etree.SubElement(pPr, f"{{{W_NS}}}bidi")
            bidi.set(f"{{{W_NS}}}val", "1")
            rPr = run._r.get_or_add_rPr()
            for existing in rPr.findall(f"{{{W_NS}}}rtl"):
                rPr.remove(existing)
            rtl = etree.SubElement(rPr, f"{{{W_NS}}}rtl")
            rtl.set(f"{{{W_NS}}}val", "1")
            p.paragraph_format.space_after = Pt(style.paragraph_spacing_pt)
        else:
            _apply_rtl_to_paragraph(
                doc.add_paragraph(text),
                font_size_pt=style.font_size_pt,
                space_after_pt=style.paragraph_spacing_pt,
            )

    # Closing phrase
    if style.include_closing_phrase:
        doc.add_paragraph("")  # spacing
        p = doc.add_paragraph(style.closing_phrase)
        _apply_rtl_to_paragraph(p, font_size_pt=style.font_size_pt,
                                space_after_pt=style.paragraph_spacing_pt)

    # Signature block
    if style.include_signature_block:
        doc.add_paragraph("")  # spacing
        p = doc.add_paragraph(style.signature_name_placeholder)
        _apply_rtl_to_paragraph(p, font_size_pt=style.font_size_pt,
                                space_after_pt=style.paragraph_spacing_pt)
        p = doc.add_paragraph(style.signature_title)
        _apply_rtl_to_paragraph(p, font_size_pt=style.font_size_pt, bold=True,
                                space_after_pt=style.paragraph_spacing_pt)


def build_letter_from_style(
    draft: GeneratedDraft,
    style: LetterStyle | None = None,
    *,
    out_path: Path | None = None,
) -> bytes:
    """Build a fresh DOCX from ``draft`` + ``style`` — no template required.

    This is the template-independent entry point. The output is an
    official, ready-to-send letter with:
      * explicit RTL formatting (bidi + rtl + right-align) on every paragraph
      * the standard Saudi administrative letter skeleton (Basmala,
        date row, recipient block, greeting, body, closing, signature)
      * body cleaned of inline citation tags and the ``## المصادر`` section
      * configurable style (font, margins, basmala on/off, etc.)

    The returned bytes are the same shape as ``build_letter_from_template``;
    callers that previously received ``(bytes, profile)`` from the legacy
    path can use ``build_letter`` instead for a uniform interface.
    """
    try:
        from docx import Document
    except ImportError as e:  # noqa: BLE001
        raise RuntimeError("python-docx is required for DOCX export") from e

    style = style or DEFAULT_LETTER_STYLE
    doc = Document()  # blank document
    apply_letter_style(doc, style)
    _build_letter_skeleton(doc, style, draft)

    if out_path is not None:
        doc.save(str(out_path))
        with open(out_path, "rb") as f:
            return f.read()

    # Default: return bytes from an in-memory save
    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_letter(
    template_or_style: Path | LetterStyle | None,
    draft: GeneratedDraft,
    *,
    placeholders: dict[str, str] | None = None,
    out_path: Path | None = None,
) -> tuple[bytes, TemplateProfile | None]:
    """Dispatcher: route to the right builder based on the first argument.

    * ``Path`` (template file)            → legacy ``build_letter_from_template``
    * ``LetterStyle`` (or None)          → new ``build_letter_from_style``

    Returns ``(docx_bytes, template_profile_or_None)``. The profile is
    None when no template was used (the new style-driven path).
    """
    if isinstance(template_or_style, Path):
        # Legacy template-driven path (unchanged)
        return build_letter_from_template(
            template_or_style,
            draft,
            placeholders=placeholders,
            include_sources=False,
            out_path=out_path,
        )
    # LetterStyle or None → style-driven path
    style = template_or_style if isinstance(template_or_style, LetterStyle) else None
    data = build_letter_from_style(draft, style=style, out_path=out_path)
    return data, None
